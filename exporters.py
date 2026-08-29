from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass(slots=True)
class PageResult:
    page: int
    source: str
    translation: str = ""


def _clean_title(filename: str) -> str:
    return Path(filename).stem.strip() or "OCR document"


def write_markdown(
    output_path: Path,
    *,
    filename: str,
    pages: list[PageResult],
    include_source: bool,
    include_translation: bool,
    target_label: str = "",
    document_translation: str = "",
) -> None:
    lines = [f"# {_clean_title(filename)}", ""]
    if include_source:
        if include_translation:
            lines.extend(["## Original", ""])
        for result in pages:
            lines.extend([f"## Page {result.page}", ""])
            lines.extend([result.source.strip(), ""])

    if include_translation:
        lines.extend([f"## Translation{f' — {target_label}' if target_label else ''}", ""])
        if document_translation.strip():
            lines.extend([document_translation.strip(), ""])
        else:
            # Preserve the old PageResult-based API for callers that do not
            # provide a document-level translation.
            for result in pages:
                if result.translation.strip():
                    lines.extend([result.translation.strip(), ""])
    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _set_run_font(run, size: int = 11, bold: bool = False, color: str = "28231F") -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK JP")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_text_blocks(document: Document, text: str) -> None:
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    for block in blocks:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.2
        _set_run_font(paragraph.add_run(block))


def write_docx(
    output_path: Path,
    *,
    filename: str,
    pages: list[PageResult],
    include_source: bool,
    include_translation: bool,
    target_label: str = "",
    document_translation: str = "",
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = document.add_heading(level=0)
    _set_run_font(title.add_run(_clean_title(filename)), size=20, bold=True, color="9D2C21")

    if include_source:
        if include_translation:
            heading = document.add_heading(level=1)
            _set_run_font(heading.add_run("Original"), size=14, bold=True, color="9D2C21")
        for result in pages:
            page_heading = document.add_heading(level=2 if include_translation else 1)
            _set_run_font(
                page_heading.add_run(f"Page {result.page}"),
                size=12 if include_translation else 14,
                bold=True,
            )
            _add_text_blocks(document, result.source)

    if include_translation:
        heading = document.add_heading(level=1)
        label = f"Translation — {target_label}" if target_label else "Translation"
        _set_run_font(heading.add_run(label), size=14, bold=True, color="9D2C21")
        if document_translation.strip():
            _add_text_blocks(document, document_translation)
        else:
            # Preserve the old PageResult-based API for callers that do not
            # provide a document-level translation.
            for result in pages:
                if result.translation.strip():
                    _add_text_blocks(document, result.translation)

    document.save(output_path)
